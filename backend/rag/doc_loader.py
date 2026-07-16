# backend/rag/doc_loader.py
# 🆕 文档加载器：支持 PDF / DOCX 文件上传，提取文本内容。
#
# 支持 PDF/DOCX 文件上传并提取文本内容。
#
# 支持格式：
#   - PDF:  PyMuPDF (fitz) 提取文本
#   - DOCX: python-docx 提取段落 + 表格
#
# 使用方式：
#   loader = DocLoader()
#   docs = loader.load("path/to/file.pdf")  # → list[Document]

import os
from langchain_core.documents import Document
from backend.core.logger import get_logger

logger = get_logger(__name__)

# 支持的文件类型
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".md", ".markdown", ".txt"}


class DocLoader:
    """统一文档加载器 — 根据文件扩展名自动选择合适的解析器。

    解析结果返回 LangChain Document 列表，可直接喂给 chunker → 向量库。
    """

    def load(self, file_path: str) -> list[Document]:
        """加载文档文件，提取文本内容。

        :param file_path: 文件绝对路径
        :return: list[Document] — 每页/每段一个 Document
        """
        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".pdf":
            return self._load_pdf(file_path)
        elif ext == ".docx":
            return self._load_docx(file_path)
        elif ext in (".md", ".markdown", ".txt"):
            return self._load_markdown(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}，仅支持 PDF / DOCX / Markdown / TXT")

    # ── PDF 解析（PyMuPDF）───────────────────────────

    @staticmethod
    def _load_pdf(file_path: str) -> list[Document]:
        """使用 PyMuPDF 提取 PDF 文本。"""
        documents = []
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            full_text = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    full_text.append(text.strip())

            if full_text:
                content = "\n\n".join(full_text)
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source_file": os.path.basename(file_path),
                        "type": "pdf",
                        "pages": len(doc),
                        "file_path": file_path,
                    },
                ))
                logger.info("doc_loader.pdf_loaded",
                           file=os.path.basename(file_path),
                           pages=len(doc),
                           chars=len(content))
            else:
                logger.warning("doc_loader.pdf_empty", file=file_path)

            doc.close()
        except ImportError:
            logger.error("doc_loader.pdf_no_pymupdf",
                        hint="pip install PyMuPDF")
            raise
        except Exception as e:
            logger.error("doc_loader.pdf_failed", file=file_path, error=str(e))
            raise

        return documents

    # ── DOCX 解析（python-docx）───────────────────────

    @staticmethod
    def _load_docx(file_path: str) -> list[Document]:
        """使用 python-docx 提取 Word 文档的段落和表格。"""
        documents = []
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            full_text = []

            # 提取段落
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # 提取表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_text:
                        full_text.append(" | ".join(row_text))

            if full_text:
                content = "\n".join(full_text)
                documents.append(Document(
                    page_content=content,
                    metadata={
                        "source_file": os.path.basename(file_path),
                        "type": "docx",
                        "file_path": file_path,
                    },
                ))
                logger.info("doc_loader.docx_loaded",
                           file=os.path.basename(file_path),
                           chars=len(content))
            else:
                logger.warning("doc_loader.docx_empty", file=file_path)

        except ImportError:
            logger.error("doc_loader.docx_no_python_docx",
                        hint="pip install python-docx")
            raise
        except Exception as e:
            logger.error("doc_loader.docx_failed", file=file_path, error=str(e))
            raise

        return documents

    # ── Markdown / TXT 解析 ──────────────────────────

    @staticmethod
    def _load_markdown(file_path: str) -> list[Document]:
        """读取 Markdown / TXT 文件，保留原始格式。"""
        documents = []
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()

            if content.strip():
                documents.append(Document(
                    page_content=content.strip(),
                    metadata={
                        "source_file": os.path.basename(file_path),
                        "type": os.path.splitext(file_path)[1].replace(".", ""),
                        "file_path": file_path,
                    },
                ))
                logger.info("doc_loader.markdown_loaded",
                           file=os.path.basename(file_path), chars=len(content))
            else:
                logger.warning("doc_loader.markdown_empty", file=file_path)
        except UnicodeDecodeError:
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    content = f.read()
                if content.strip():
                    documents.append(Document(
                        page_content=content.strip(),
                        metadata={
                            "source_file": os.path.basename(file_path),
                            "type": os.path.splitext(file_path)[1].replace(".", ""),
                        },
                    ))
            except Exception as e:
                logger.error("doc_loader.markdown_failed", file=file_path, error=str(e))
                raise
        except Exception as e:
            logger.error("doc_loader.markdown_failed", file=file_path, error=str(e))
            raise

        return documents


# ── 测试代码 ───────────────────────────────────────────────
if __name__ == "__main__":
    loader = DocLoader()
    print("DocLoader 就绪，支持格式:", ALLOWED_EXTENSIONS)
    print("PyMuPDF:", "已安装" if __import__("importlib.util").util.find_spec("fitz") else "未安装")
    print("python-docx:", "已安装" if __import__("importlib.util").util.find_spec("docx") else "未安装")
